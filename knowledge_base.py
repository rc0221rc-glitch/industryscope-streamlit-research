from __future__ import annotations

import hashlib
import json
import math
import os
import re
import shutil
import tempfile
from dataclasses import dataclass, asdict
from datetime import datetime
from io import BytesIO
from pathlib import Path
from typing import Any
from zipfile import ZIP_DEFLATED, ZipFile

from bs4 import BeautifulSoup

try:
    import fitz  # PyMuPDF
except ImportError:
    fitz = None

try:
    import docx
except ImportError:
    docx = None

try:
    import pandas as pd
except ImportError:
    pd = None

try:
    import streamlit as st
except ImportError:
    st = None


KB_ROOT = Path(os.getenv("INDUSTRYSCOPE_KB_DIR", "data/knowledge_base"))
KB_FILES_DIR = KB_ROOT / "files"
KB_INDEX_DIR = KB_ROOT / "index"
KB_CHUNKS_PATH = KB_INDEX_DIR / "chunks.jsonl"
KB_DOCS_PATH = KB_INDEX_DIR / "documents.json"
KB_SNAPSHOT_ENV_KEYS = {
    "bucket": "INDUSTRYSCOPE_KB_S3_BUCKET",
    "key": "INDUSTRYSCOPE_KB_S3_KEY",
    "endpoint": "INDUSTRYSCOPE_KB_S3_ENDPOINT_URL",
    "region": "INDUSTRYSCOPE_KB_S3_REGION",
    "access_key": "INDUSTRYSCOPE_KB_S3_ACCESS_KEY_ID",
    "secret_key": "INDUSTRYSCOPE_KB_S3_SECRET_ACCESS_KEY",
}


SOURCE_TYPE_TIERS = {
    "公司公告/年报/招股书": "T0",
    "专利报告/专利地图": "T1",
    "论文/会议": "T0",
    "专家纪要/调研纪要": "T1",
    "券商/投行/咨询研报": "T1",
    "行业白皮书/协会报告": "T1",
    "内部笔记": "T2",
    "公众号/媒体转载": "T3",
    "微信公众号候选线索": "T3",
    "其他": "T2",
}


WECHAT_CANDIDATE_SOURCE_TYPE = "微信公众号候选线索"


def is_wechat_candidate_stub_record(item: dict[str, Any]) -> bool:
    source_type = str(item.get("source_type", ""))
    title = str(item.get("title", ""))
    text = str(item.get("text", "") or item.get("snippet", ""))
    return (
        source_type == WECHAT_CANDIDATE_SOURCE_TYPE
        or "候选线索" in title
        or "入库状态：候选线索" in text
        or "未能自动抓取全文" in text
    )


@dataclass
class KBDocument:
    doc_id: str
    title: str
    filename: str
    stored_path: str
    source_type: str
    source_tier: str
    source_org: str
    publish_date: str
    industry_tags: str
    company_tags: str
    technology_tags: str
    created_at: str
    chunk_count: int
    source_url: str = ""


@dataclass
class KBChunk:
    chunk_id: str
    doc_id: str
    title: str
    filename: str
    source_type: str
    source_tier: str
    source_org: str
    publish_date: str
    industry_tags: str
    company_tags: str
    technology_tags: str
    page: str
    section: str
    text: str
    token_hint: int
    source_url: str = ""


def ensure_kb_dirs() -> None:
    KB_FILES_DIR.mkdir(parents=True, exist_ok=True)
    KB_INDEX_DIR.mkdir(parents=True, exist_ok=True)
    if not KB_DOCS_PATH.exists():
        KB_DOCS_PATH.write_text("{}", encoding="utf-8")
    if not KB_CHUNKS_PATH.exists():
        KB_CHUNKS_PATH.write_text("", encoding="utf-8")


def filename_safe(value: str) -> str:
    safe = re.sub(r"[^\w\u4e00-\u9fff.-]+", "_", value.strip())
    return safe.strip("_") or "document"


def file_sha256(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as handle:
        for block in iter(lambda: handle.read(1024 * 1024), b""):
            digest.update(block)
    return digest.hexdigest()


def normalize_text(text: str) -> str:
    text = (text or "").replace("\xa0", " ")
    text = re.sub(r"[ \t\r\f\v]+", " ", text)
    text = re.sub(r" *\n *", "\n", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    lines: list[str] = []
    prev = ""
    for raw in text.splitlines():
        line = raw.strip()
        if not line or line == prev:
            continue
        lines.append(line)
        prev = line
    return "\n".join(lines).strip()


def extract_pdf(path: Path) -> list[dict[str, str]]:
    if fitz is None:
        raise RuntimeError("PyMuPDF 未安装，无法解析 PDF。")
    pages: list[dict[str, str]] = []
    with fitz.open(path) as doc:
        for index, page in enumerate(doc, start=1):
            text = normalize_text(page.get_text("text"))
            if text:
                pages.append({"page": str(index), "section": "", "text": text})
    return pages


def extract_docx(path: Path) -> list[dict[str, str]]:
    if docx is None:
        raise RuntimeError("python-docx 未安装，无法解析 DOCX。")
    document = docx.Document(str(path))
    parts: list[str] = []
    current_heading = ""
    sections: list[dict[str, str]] = []
    for para in document.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        style = (para.style.name or "").lower() if para.style else ""
        if "heading" in style and parts:
            sections.append({"page": "", "section": current_heading, "text": normalize_text("\n".join(parts))})
            parts = []
            current_heading = text
        elif "heading" in style:
            current_heading = text
        else:
            parts.append(text)
    if parts:
        sections.append({"page": "", "section": current_heading, "text": normalize_text("\n".join(parts))})
    return sections or [{"page": "", "section": "", "text": normalize_text("\n".join(p.text for p in document.paragraphs))}]


def extract_spreadsheet(path: Path) -> list[dict[str, str]]:
    if pd is None:
        raise RuntimeError("pandas/openpyxl 未安装，无法解析表格。")
    sheets = pd.read_excel(path, sheet_name=None) if path.suffix.lower() in {".xlsx", ".xls"} else {"csv": pd.read_csv(path)}
    sections: list[dict[str, str]] = []
    for sheet_name, frame in sheets.items():
        text = frame.fillna("").astype(str).to_csv(index=False)
        text = normalize_text(text)
        if text:
            sections.append({"page": "", "section": str(sheet_name), "text": text})
    return sections


def extract_html(path: Path) -> list[dict[str, str]]:
    soup = BeautifulSoup(path.read_text(encoding="utf-8", errors="ignore"), "html.parser")
    for tag in soup(["script", "style", "nav", "footer", "header", "aside", "form"]):
        tag.decompose()
    return [{"page": "", "section": "", "text": normalize_text(soup.get_text("\n", strip=True))}]


def extract_plain(path: Path) -> list[dict[str, str]]:
    return [{"page": "", "section": "", "text": normalize_text(path.read_text(encoding="utf-8", errors="ignore"))}]


def extract_document(path: Path) -> list[dict[str, str]]:
    suffix = path.suffix.lower()
    if suffix == ".pdf":
        return extract_pdf(path)
    if suffix == ".docx":
        return extract_docx(path)
    if suffix in {".xlsx", ".xls", ".csv"}:
        return extract_spreadsheet(path)
    if suffix in {".html", ".htm"}:
        return extract_html(path)
    if suffix in {".md", ".txt"}:
        return extract_plain(path)
    raise RuntimeError(f"暂不支持的文件类型：{suffix}")


def split_chunk_text(text: str, target_chars: int = 1100, overlap_chars: int = 160) -> list[str]:
    text = normalize_text(text)
    if not text:
        return []
    paragraphs = re.split(r"\n{2,}", text)
    chunks: list[str] = []
    current = ""
    for para in paragraphs:
        para = para.strip()
        if not para:
            continue
        if len(current) + len(para) + 2 <= target_chars:
            current = f"{current}\n\n{para}".strip()
        else:
            if current:
                chunks.append(current)
            if len(para) <= target_chars:
                current = para
            else:
                start = 0
                while start < len(para):
                    chunks.append(para[start:start + target_chars])
                    start += max(1, target_chars - overlap_chars)
                current = ""
    if current:
        chunks.append(current)
    if overlap_chars > 0 and len(chunks) > 1:
        overlapped: list[str] = []
        previous_tail = ""
        for chunk in chunks:
            combined = f"{previous_tail}\n{chunk}".strip() if previous_tail else chunk
            overlapped.append(combined[: target_chars + overlap_chars])
            previous_tail = chunk[-overlap_chars:]
        chunks = overlapped
    return chunks


def load_documents() -> dict[str, dict[str, Any]]:
    ensure_kb_dirs()
    try:
        return json.loads(KB_DOCS_PATH.read_text(encoding="utf-8") or "{}")
    except json.JSONDecodeError:
        return {}


def save_documents(docs: dict[str, dict[str, Any]]) -> None:
    ensure_kb_dirs()
    KB_DOCS_PATH.write_text(json.dumps(docs, ensure_ascii=False, indent=2), encoding="utf-8")


def load_chunks() -> list[dict[str, Any]]:
    ensure_kb_dirs()
    chunks: list[dict[str, Any]] = []
    for line in KB_CHUNKS_PATH.read_text(encoding="utf-8").splitlines():
        if line.strip():
            try:
                chunks.append(json.loads(line))
            except json.JSONDecodeError:
                continue
    return chunks


def save_chunks(chunks: list[dict[str, Any]]) -> None:
    ensure_kb_dirs()
    KB_CHUNKS_PATH.write_text("\n".join(json.dumps(chunk, ensure_ascii=False) for chunk in chunks), encoding="utf-8")


def kb_file_size_bytes() -> int:
    ensure_kb_dirs()
    total = 0
    for path in KB_ROOT.rglob("*"):
        if path.is_file():
            try:
                total += path.stat().st_size
            except OSError:
                continue
    return total


def format_bytes(value: int) -> str:
    size = float(value)
    for unit in ["B", "KB", "MB", "GB", "TB"]:
        if size < 1024 or unit == "TB":
            return f"{size:.1f} {unit}" if unit != "B" else f"{int(size)} B"
        size /= 1024
    return f"{value} B"


def export_kb_snapshot_bytes() -> bytes:
    ensure_kb_dirs()
    buffer = BytesIO()
    with ZipFile(buffer, "w", compression=ZIP_DEFLATED) as archive:
        manifest = {
            "created_at": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
            "root": str(KB_ROOT),
            "documents": len(load_documents()),
            "chunks": len(load_chunks()),
            "schema": "industryscope-kb-snapshot-v1",
        }
        archive.writestr("manifest.json", json.dumps(manifest, ensure_ascii=False, indent=2).encode("utf-8"))
        archive.write(KB_DOCS_PATH, "index/documents.json")
        archive.write(KB_CHUNKS_PATH, "index/chunks.jsonl")
        if KB_FILES_DIR.exists():
            for file_path in KB_FILES_DIR.rglob("*"):
                if file_path.is_file():
                    archive.write(file_path, f"files/{file_path.relative_to(KB_FILES_DIR).as_posix()}")
    return buffer.getvalue()


def import_kb_snapshot_bytes(snapshot: bytes, merge: bool = True) -> dict[str, Any]:
    ensure_kb_dirs()
    with tempfile.TemporaryDirectory() as tmpdir:
        tmp_root = Path(tmpdir)
        with ZipFile(BytesIO(snapshot)) as archive:
            for info in archive.infolist():
                target = (tmp_root / info.filename).resolve()
                if not str(target).startswith(str(tmp_root.resolve())):
                    raise RuntimeError("知识库快照包含不安全路径，已拒绝导入。")
            archive.extractall(tmp_root)

        docs_path = tmp_root / "index" / "documents.json"
        chunks_path = tmp_root / "index" / "chunks.jsonl"
        files_dir = tmp_root / "files"
        if not docs_path.exists() or not chunks_path.exists():
            raise RuntimeError("知识库快照缺少 index/documents.json 或 index/chunks.jsonl。")

        incoming_docs = json.loads(docs_path.read_text(encoding="utf-8") or "{}")
        incoming_chunks: list[dict[str, Any]] = []
        for line in chunks_path.read_text(encoding="utf-8").splitlines():
            if line.strip():
                incoming_chunks.append(json.loads(line))

        if merge:
            docs = load_documents()
            docs.update(incoming_docs)
            existing_doc_ids = {chunk.get("doc_id") for chunk in incoming_chunks}
            chunks = [chunk for chunk in load_chunks() if chunk.get("doc_id") not in existing_doc_ids]
            chunks.extend(incoming_chunks)
        else:
            docs = incoming_docs
            chunks = incoming_chunks
            if KB_FILES_DIR.exists():
                shutil.rmtree(KB_FILES_DIR)
            KB_FILES_DIR.mkdir(parents=True, exist_ok=True)

        if files_dir.exists():
            for file_path in files_dir.rglob("*"):
                if file_path.is_file():
                    target = KB_FILES_DIR / file_path.relative_to(files_dir)
                    target.parent.mkdir(parents=True, exist_ok=True)
                    shutil.copyfile(file_path, target)

        save_documents(docs)
        save_chunks(chunks)
        return {"documents": len(incoming_docs), "chunks": len(incoming_chunks), "merge": merge}


def s3_config() -> dict[str, str]:
    values: dict[str, str] = {}
    for name, env_key in KB_SNAPSHOT_ENV_KEYS.items():
        value = ""
        if st is not None:
            try:
                value = str(st.secrets.get(env_key, "") or "").strip()
            except Exception:
                value = ""
        values[name] = value or os.getenv(env_key, "").strip()
    return values


def s3_sync_enabled() -> bool:
    cfg = s3_config()
    return bool(cfg["bucket"] and cfg["access_key"] and cfg["secret_key"])


def _boto3_client() -> Any:
    try:
        import boto3
    except ImportError as exc:
        raise RuntimeError("未安装 boto3，无法使用 S3/R2 知识库持久化。请在 requirements.txt 中加入 boto3。") from exc
    cfg = s3_config()
    kwargs: dict[str, Any] = {
        "aws_access_key_id": cfg["access_key"],
        "aws_secret_access_key": cfg["secret_key"],
    }
    if cfg["endpoint"]:
        kwargs["endpoint_url"] = cfg["endpoint"]
    if cfg["region"]:
        kwargs["region_name"] = cfg["region"]
    return boto3.client("s3", **kwargs)


def upload_kb_snapshot_to_s3() -> str:
    cfg = s3_config()
    if not s3_sync_enabled():
        return "未配置远端知识库持久化，跳过上传。"
    key = cfg["key"] or "industryscope/kb_snapshot.zip"
    client = _boto3_client()
    client.put_object(
        Bucket=cfg["bucket"],
        Key=key,
        Body=export_kb_snapshot_bytes(),
        ContentType="application/zip",
    )
    return f"已上传知识库快照到 s3://{cfg['bucket']}/{key}"


def restore_kb_snapshot_from_s3(merge: bool = True) -> dict[str, Any]:
    cfg = s3_config()
    if not s3_sync_enabled():
        raise RuntimeError("未配置远端知识库持久化。")
    key = cfg["key"] or "industryscope/kb_snapshot.zip"
    client = _boto3_client()
    response = client.get_object(Bucket=cfg["bucket"], Key=key)
    data = response["Body"].read()
    result = import_kb_snapshot_bytes(data, merge=merge)
    result["remote"] = f"s3://{cfg['bucket']}/{key}"
    return result


def try_restore_from_s3_if_empty() -> str:
    if not s3_sync_enabled():
        return ""
    if load_documents() or load_chunks():
        return ""
    try:
        result = restore_kb_snapshot_from_s3(merge=False)
        return f"已从远端恢复知识库：{result.get('documents', 0)} 个文档，{result.get('chunks', 0)} 个片段。"
    except Exception as exc:
        return f"远端知识库自动恢复失败：{exc}"


def autosync_kb_snapshot() -> str:
    if not s3_sync_enabled():
        return ""
    try:
        return upload_kb_snapshot_to_s3()
    except Exception as exc:
        return f"知识库远端同步失败：{exc}"


def add_document(
    source_path: Path,
    title: str = "",
    source_type: str = "其他",
    source_org: str = "",
    publish_date: str = "",
    industry_tags: str = "",
    company_tags: str = "",
    technology_tags: str = "",
    source_url: str = "",
) -> KBDocument:
    ensure_kb_dirs()
    digest = file_sha256(source_path)
    doc_id = digest[:16]
    suffix = source_path.suffix.lower()
    stored_name = f"{doc_id}_{filename_safe(source_path.stem)}{suffix}"
    stored_path = KB_FILES_DIR / stored_name
    if source_path.resolve() != stored_path.resolve():
        shutil.copyfile(source_path, stored_path)

    docs = load_documents()
    existing_chunks = [chunk for chunk in load_chunks() if chunk.get("doc_id") != doc_id]
    sections = extract_document(stored_path)
    source_tier = SOURCE_TYPE_TIERS.get(source_type, "T2")
    doc_title = title.strip() or source_path.stem
    chunks: list[dict[str, Any]] = []
    chunk_index = 0
    for section in sections:
        for text in split_chunk_text(section.get("text", "")):
            chunk_index += 1
            chunk_id = f"{doc_id}-{chunk_index:04d}"
            chunk = KBChunk(
                chunk_id=chunk_id,
                doc_id=doc_id,
                title=doc_title,
                filename=stored_name,
                source_type=source_type,
                source_tier=source_tier,
                source_org=source_org,
                publish_date=publish_date,
                industry_tags=industry_tags,
                company_tags=company_tags,
                technology_tags=technology_tags,
                page=section.get("page", ""),
                section=section.get("section", ""),
                text=text,
                token_hint=max(1, len(text) // 2),
                source_url=source_url,
            )
            chunks.append(asdict(chunk))

    document = KBDocument(
        doc_id=doc_id,
        title=doc_title,
        filename=stored_name,
        stored_path=str(stored_path),
        source_type=source_type,
        source_tier=source_tier,
        source_org=source_org,
        publish_date=publish_date,
        industry_tags=industry_tags,
        company_tags=company_tags,
        technology_tags=technology_tags,
        created_at=datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
        chunk_count=len(chunks),
        source_url=source_url,
    )
    docs[doc_id] = asdict(document)
    save_documents(docs)
    save_chunks(existing_chunks + chunks)
    return document


def delete_document(doc_id: str) -> None:
    docs = load_documents()
    doc = docs.pop(doc_id, None)
    if doc:
        try:
            Path(doc.get("stored_path", "")).unlink(missing_ok=True)
        except Exception:
            pass
    save_documents(docs)
    save_chunks([chunk for chunk in load_chunks() if chunk.get("doc_id") != doc_id])


def tokenize_query(text: str) -> list[str]:
    text = (text or "").lower()
    tokens = re.findall(r"[\u4e00-\u9fff]{2,}|[a-zA-Z][a-zA-Z0-9\-]{2,}", text)
    expanded: list[str] = []
    for token in tokens:
        expanded.append(token)
        if re.fullmatch(r"[\u4e00-\u9fff]{4,}", token):
            expanded.extend(token[i:i + 2] for i in range(len(token) - 1))
    stop = {"行业", "研究", "报告", "市场", "分析", "the", "and", "for", "with"}
    return [token for token in expanded if token not in stop]


def kb_core_terms(query: str) -> set[str]:
    text = (query or "").lower()
    if any(term in text for term in ["硅光", "光芯片", "silicon photonic", "silicon photonics"]):
        return {
            "硅光", "硅光芯片", "硅光子", "光芯片", "光子集成", "光通信", "光模块",
            "cpo", "co-packaged", "copackaged", "photonic", "photonics",
            "silicon photonics", "optical transceiver",
        }
    if any(term in text for term in ["脑机", "bci", "brain-computer", "neural interface"]):
        return {"脑机", "脑机接口", "bci", "brain-computer", "neuralink", "神经接口"}
    if any(term in text for term in ["肌电", "emg", "semg"]):
        return {"肌电", "emg", "semg", "electromyography", "电极", "腕带", "手环"}
    if any(term in text for term in ["玻璃基板", "glass substrate", "glass core"]):
        return {"玻璃基板", "玻璃通孔", "tgv", "glass substrate", "glass core", "advanced packaging"}

    generic = {
        "行业", "产业", "领域", "赛道", "研究", "报告", "市场", "分析", "技术", "路线",
        "最新", "进展", "融资", "客户", "产线", "专利", "厂商", "公司", "国内外",
        "商业化", "供应链", "成熟度", "上游", "the", "and", "for", "with", "market",
        "industry", "technology", "company", "customer", "patent", "report",
    }
    terms = {term for term in tokenize_query(query) if term not in generic and len(term) >= 2}
    noisy_fragments = {"技术", "路线", "最新", "进展", "融资", "客户", "产线", "专利", "厂商", "公司", "国内", "内外", "外厂", "术路", "新进", "户产", "业拐", "拐点", "商化", "条件", "供应", "应链", "成熟", "熟度", "上游"}
    return {term for term in terms if term not in noisy_fragments}


def kb_matches_core_terms(haystack: str, core_terms: set[str]) -> bool:
    if not core_terms:
        return True
    return any(term.lower() in haystack for term in core_terms)


def date_score(value: str) -> float:
    match = re.search(r"(20\d{2})", value or "")
    if not match:
        return 0.0
    year = int(match.group(1))
    return max(0.0, min(6.0, (year - 2020) * 1.2))


def tier_score(tier: str) -> float:
    return {"T0": 16.0, "T1": 12.0, "T2": 7.0, "T3": 2.0}.get(tier, 4.0)


def search_knowledge_base(query: str, top_k: int = 12, filters: dict[str, str] | None = None) -> list[dict[str, Any]]:
    filters = filters or {}
    query_tokens = tokenize_query(query)
    if not query_tokens:
        return []
    query_terms = set(query_tokens)
    core_terms = kb_core_terms(query)
    chunks = load_chunks()
    scored: list[dict[str, Any]] = []
    for chunk in chunks:
        haystack = " ".join(
            str(chunk.get(key, ""))
            for key in ["title", "source_org", "industry_tags", "company_tags", "technology_tags", "section", "text"]
        ).lower()
        tag_haystack = f"{chunk.get('industry_tags', '')} {chunk.get('company_tags', '')} {chunk.get('technology_tags', '')}".lower()
        if filters.get("source_tier") and chunk.get("source_tier") != filters["source_tier"]:
            continue
        if filters.get("source_type") and chunk.get("source_type") != filters["source_type"]:
            continue
        if not kb_matches_core_terms(haystack, core_terms):
            continue
        hits = 0
        dense_hits = 0
        for term in query_terms:
            count = haystack.count(term.lower())
            if count:
                hits += 1
                dense_hits += min(count, 5)
        if hits == 0:
            continue
        coverage = hits / max(1, len(query_terms))
        phrase_bonus = 10.0 if query.lower() in haystack else 0.0
        title_bonus = 5.0 if any(term in str(chunk.get("title", "")).lower() for term in query_terms) else 0.0
        tag_bonus = 14.0 if any(term in tag_haystack for term in query_terms) else 0.0
        core_bonus = 18.0 if any(term.lower() in tag_haystack for term in core_terms) else 6.0
        score = coverage * 38 + math.log1p(dense_hits) * 8 + tier_score(str(chunk.get("source_tier", ""))) + date_score(str(chunk.get("publish_date", ""))) + phrase_bonus + title_bonus + tag_bonus + core_bonus
        if is_wechat_candidate_stub_record(chunk):
            score = max(1.0, score - 25.0)
        result = dict(chunk)
        result["score"] = round(score, 2)
        result["match_coverage"] = round(coverage, 2)
        result["matched_terms"] = ", ".join(sorted(term for term in query_terms if term in haystack)[:12])
        result["core_terms"] = ", ".join(sorted(core_terms)[:12])
        scored.append(result)
    scored.sort(key=lambda item: item["score"], reverse=True)
    return scored[:top_k]


def kb_stats() -> dict[str, Any]:
    docs = load_documents()
    chunks = load_chunks()
    size_bytes = kb_file_size_bytes()
    type_counts: dict[str, int] = {}
    tier_counts: dict[str, int] = {}
    for doc in docs.values():
        type_counts[doc.get("source_type", "其他")] = type_counts.get(doc.get("source_type", "其他"), 0) + 1
        tier_counts[doc.get("source_tier", "T2")] = tier_counts.get(doc.get("source_tier", "T2"), 0) + 1
    return {
        "documents": len(docs),
        "chunks": len(chunks),
        "type_counts": type_counts,
        "tier_counts": tier_counts,
        "root": str(KB_ROOT),
        "size_bytes": size_bytes,
        "size_human": format_bytes(size_bytes),
        "remote_sync": s3_sync_enabled(),
    }


def export_kb_json() -> str:
    return json.dumps({"documents": load_documents(), "chunks": load_chunks()}, ensure_ascii=False, indent=2)


def kb_results_to_sources(results: list[dict[str, Any]]) -> list[dict[str, str]]:
    sources: list[dict[str, str]] = []
    for idx, item in enumerate(results, start=1):
        page = f"p.{item.get('page')}" if item.get("page") else item.get("section", "")
        title = item.get("title") or item.get("filename") or f"KB Source {idx}"
        label = f"{title} {page}".strip()
        chunk_id = item.get("chunk_id", "")
        is_candidate_stub = is_wechat_candidate_stub_record(item)
        source_channel = "微信公众号候选线索" if is_candidate_stub else "专属知识库"
        use_policy = (
            "仅可作为待补全文的发现线索；不得作为已读全文证据，不得支撑市场规模、份额、融资、订单、财务、客户绑定、技术指标或全球第一等强结论。"
            if is_candidate_stub
            else "可作为私有知识库证据；重大强结论仍需与一手公开来源或其他知识库文件交叉验证。"
        )
        sources.append({
            "title": label,
            "url": f"kb://{chunk_id}",
            "snippet": str(item.get("text", ""))[:500],
            "type": "knowledge_base",
            "source_tier": item.get("source_tier", "T2"),
            "source_channel": source_channel,
            "density_band": "候选线索" if is_candidate_stub else ("高" if float(item.get("score", 0)) >= 55 else "中"),
            "evidence_density": str(int(float(item.get("score", 0)))),
            "relevance": str(int(float(item.get("score", 0)) // 10)),
            "use_policy": use_policy,
            "kb_is_candidate_stub": "true" if is_candidate_stub else "false",
            "kb_chunk_id": chunk_id,
            "kb_doc_id": item.get("doc_id", ""),
            "kb_filename": item.get("filename", ""),
            "kb_page": item.get("page", ""),
            "kb_section": item.get("section", ""),
            "source_url": item.get("source_url", ""),
        })
    return sources
